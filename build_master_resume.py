from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("/Users/sanjanasawant/Documents/Website/Sanjana_Sawant_Master_Resume.docx")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, position="bottom", color="DADCE0", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{position}"))
    if border is None:
        border = OxmlElement(f"w:{position}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Title", "Subtitle"):
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Arial"

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x11, 0x1B, 0x21)
    title.paragraph_format.space_after = Pt(2)

    if "ResumeHeading" not in doc.styles:
        heading = doc.styles.add_style("ResumeHeading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        heading = doc.styles["ResumeHeading"]
    heading.font.name = "Arial"
    heading.font.size = Pt(10.3)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(0x17, 0x33, 0x49)
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(2)
    heading.paragraph_format.keep_with_next = True

    if "EntryHeader" not in doc.styles:
        entry_header = doc.styles.add_style("EntryHeader", WD_STYLE_TYPE.PARAGRAPH)
    else:
        entry_header = doc.styles["EntryHeader"]
    entry_header.font.name = "Arial"
    entry_header.font.size = Pt(10.2)
    entry_header.font.bold = True
    entry_header.paragraph_format.space_before = Pt(3)
    entry_header.paragraph_format.space_after = Pt(0)
    entry_header.paragraph_format.keep_with_next = True

    if "EntrySub" not in doc.styles:
        entry_sub = doc.styles.add_style("EntrySub", WD_STYLE_TYPE.PARAGRAPH)
    else:
        entry_sub = doc.styles["EntrySub"]
    entry_sub.font.name = "Arial"
    entry_sub.font.size = Pt(9.4)
    entry_sub.font.italic = True
    entry_sub.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    entry_sub.paragraph_format.space_after = Pt(1)
    entry_sub.paragraph_format.keep_with_next = True

    if "CompactBullet" not in doc.styles:
        bullet = doc.styles.add_style("CompactBullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = doc.styles["CompactBullet"]
    bullet.base_style = doc.styles["Normal"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(9.7)
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)
    bullet.paragraph_format.space_after = Pt(0)
    bullet.paragraph_format.line_spacing = 1.03

    if "SkillLine" not in doc.styles:
        skill = doc.styles.add_style("SkillLine", WD_STYLE_TYPE.PARAGRAPH)
    else:
        skill = doc.styles["SkillLine"]
    skill.base_style = doc.styles["Normal"]
    skill.font.name = "Arial"
    skill.font.size = Pt(9.8)
    skill.paragraph_format.space_after = Pt(1)


def add_contact_block(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SANJANA UMESH SAWANT")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(3)
    run = contact.add_run(
        "New Brunswick, NJ | sanjanaumesh.sawant@rutgers.edu | LinkedIn: Sanjana Umesh Sawant | GitHub: sanjanaumeshsawant0810"
    )
    run.font.name = "Arial"
    run.font.size = Pt(9.7)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    set_paragraph_border(contact)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="ResumeHeading")
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="CompactBullet")
    p.add_run(f"• {text}")


def add_two_col_entry(doc: Document, left: str, right: str, style: str = "EntryHeader") -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    widths = (Inches(5.6), Inches(1.2))
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    left_cell, right_cell = table.rows[0].cells
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    lp = left_cell.paragraphs[0]
    lp.style = doc.styles[style]
    lp.paragraph_format.space_after = Pt(0)
    lp.add_run(left)
    rp = right_cell.paragraphs[0]
    rp.style = doc.styles[style]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rp.add_run(right)


def add_summary(doc: Document) -> None:
    add_section_heading(doc, "SUMMARY")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(
        "Data science graduate building practical data and AI products across analytics, KPI reporting, ETL, applied modeling, ranking systems, and product-facing experimentation. "
        "Strongest at turning messy data into usable insight, building recruiter-friendly and stakeholder-ready reporting, and translating technical work into clear business value. "
        "Targeting data analyst, data scientist, AI, analytics engineering, and fintech-oriented roles."
    )


def add_education(doc: Document) -> None:
    add_section_heading(doc, "EDUCATION")
    add_two_col_entry(doc, "Rutgers University, M.S. in Data Science, High Honor", "May 2026")
    sub = doc.add_paragraph(style="EntrySub")
    sub.add_run("GPA: 3.6/4.0 | Coursework: Machine Learning, Regression and Time Series, DBMS, Deep Learning, LLMs, Agentic AI")
    add_two_col_entry(doc, "SRM Institute of Science and Technology, B.Tech. in Electronics and Computer Engineering", "Jun 2024")
    sub = doc.add_paragraph(style="EntrySub")
    sub.add_run("GPA: 9.01/10 | Coursework: Data Structures, Object-Oriented Design, DBMS, Embedded Systems, Computer Architecture")


def add_experience(doc: Document) -> None:
    add_section_heading(doc, "EXPERIENCE")

    add_two_col_entry(doc, "Helius Health AI | Data Science Intern | Remote", "Spring 2026")
    add_bullet(doc, "Enabled production loading of validated healthcare workforce data by building and testing local ETL/ELT workflows that staged raw survey files into Supabase/PostgreSQL and were reviewed by the CTO before handoff.")
    add_bullet(doc, "Gave sponsors visibility into onboarding, assessment completion, engagement, retention, conversion, drop-off, and specialty matching by querying Supabase/PostgreSQL through API access and turning outputs into KPI analyses.")
    add_bullet(doc, "Made one-run monitoring of user and workforce activity possible by building an end-to-end KPI reporting workflow with graphical outputs for sponsor-facing engagement and onboarding metrics.")
    add_bullet(doc, "Advanced a nursing specialty-matching decision-support model toward review by developing an initial predictive workflow that estimates which department may best fit each nursing student.")

    add_two_col_entry(doc, "NeuralSeek | AI Agent Developer Intern | Remote", "Spring 2026")
    add_bullet(doc, "Developed an AI-powered restaurant recommendation workflow using retrieval-based reasoning and recommendation-ranking logic.")
    add_bullet(doc, "Analyzed user feedback and recommendation behavior to identify factors influencing trust, relevance, and decision quality in AI-assisted restaurant discovery.")
    add_bullet(doc, "Validated the product with users and increased willingness to trust AI-generated recommendations from 17% to 78%, then extended the work into InstaDine as a fuller end-to-end application.")

    add_two_col_entry(doc, "Superpack Packaging Machines Pvt. Ltd. | PLC Programming Engineering Intern | Hyderabad, India", "Oct 2023 - Jul 2024")
    add_bullet(doc, "Improved industrial filling-machine control logic for fluids with different viscosity levels through testing and engineering iteration.")
    add_bullet(doc, "Increased machine efficiency by 4% to 8%, strengthening systems-debugging and process-optimization experience.")


def add_project(doc: Document, title: str, timeframe: str, bullets: list[str]) -> None:
    add_two_col_entry(doc, title, timeframe)
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_projects(doc: Document) -> None:
    add_section_heading(doc, "PROJECT BANK")

    add_project(
        doc,
        "InstaDine | Full-Stack AI Restaurant Recommendation Platform",
        "Spring 2026",
        [
            "Built a full-stack AI restaurant recommendation product using Flask, PostgreSQL, Docker, Google Places API, and Vertex AI/Gemini to deliver personalized, live-data-grounded recommendations.",
            "Designed conversational retrieval, ranking, and filtering workflows that combined user preferences, travel constraints, ratings, and restaurant availability to improve recommendation quality.",
            "Built session-aware chat, authentication, and recommendation persistence to turn an AI workflow into a stronger end-to-end product experience.",
        ],
    )

    add_project(
        doc,
        "Loan Default Prediction | Credit Risk Modeling",
        "2026",
        [
            "Built an end-to-end credit risk modeling project to predict loan default from 307K+ borrower records using Python, pandas, scikit-learn, and XGBoost.",
            "Compared logistic regression and XGBoost using ROC-AUC, PR-AUC, precision/recall, calibration, and threshold optimization rather than relying on accuracy alone.",
            "Connected model evaluation to lending tradeoffs by analyzing the balance between catching defaulters and avoiding unnecessary loan rejections.",
        ],
    )

    add_project(
        doc,
        "Bankruptcy Prediction Pipeline and Explainer App",
        "2026",
        [
            "Built an end-to-end bankruptcy prediction project using Python, pandas, and scikit-learn to preprocess financial data, train an interpretable classifier, and evaluate rare-event performance with ROC-AUC and confusion-matrix metrics.",
            "Implemented reusable preprocessing and modeling pipelines for scaling, encoding, missing-value imputation, and artifact persistence using joblib.",
            "Framed the workflow through a business-friendly frontend that explained why naive accuracy is misleading for imbalanced financial-risk prediction.",
        ],
    )

    add_project(
        doc,
        "Quantifying the Environmental Cost of AI",
        "Fall 2025",
        [
            "Evaluated transformer fine-tuning tradeoffs across DistilBERT, BERT-family models, and GPT-style baselines by comparing F1 performance against carbon-emissions cost on question-answering tasks.",
            "Analyzed full fine-tuning, LoRA, and more efficient training setups to quantify how emissions could drop by roughly 78% with limited performance loss.",
        ],
    )

    add_project(
        doc,
        "Effect of Technology Use on Mental Health",
        "Spring 2025",
        [
            "Processed NTIA and NSDUH survey data across multiple years, including difficult 200+ page codebooks and inconsistent historical formats.",
            "Built R and Databricks transformation workflows to standardize technology-use indicators across changing survey structures.",
            "Identified a trend in which women aged 20 to 35 reported roughly 10% higher levels of mental health distress with increased technology usage, while clearly separating correlation from causation.",
        ],
    )

    add_project(
        doc,
        "The Causal Effect of Growth Mindset Intervention on Achievement",
        "Spring 2025",
        [
            "Measured the impact of a growth mindset intervention on student performance and found a 41.2% improvement in average outcomes.",
            "Validated results using 1,000+ bootstrap samples and a 95% confidence interval, then identified that individual student expectations mattered more than school-level factors.",
        ],
    )

    add_project(
        doc,
        "Hospital Management System",
        "Fall 2025",
        [
            "Built a hospital management system inspired by Epic using Python, SQL, and SQLite for appointments, patient and doctor records, prescriptions, billing, and outcomes tracking.",
            "Implemented role-based access control to improve confidentiality and data access across administrators, doctors, and patients.",
        ],
    )

    add_project(
        doc,
        "Circuit-Vision | Autonomous Vehicle System",
        "Fall 2023",
        [
            "Worked on a prototype autonomous vehicle system involving lane detection, steering adjustment, and software performance improvements for real-time navigation.",
            "Supported hardware and software integration for a camera-based driving workflow aimed at improving vehicle control and monitoring.",
        ],
    )


def add_research_and_skills(doc: Document) -> None:
    add_section_heading(doc, "RESEARCH AND OPTIONAL SIGNALS")
    add_bullet(doc, "Published a research paper on real-time security monitoring systems at ICICIT 2024 with placement in the top 12% of submissions.")
    add_bullet(doc, "Leadership can be added back selectively for roles that value campus leadership, event coordination, or communication-heavy experience.")

    add_section_heading(doc, "SKILLS")
    lines = [
        ("Languages", "Python, SQL, R, C++, Swift, MATLAB"),
        ("Data and Analytics", "PostgreSQL, Supabase, SQLite, SQL Server, BigQuery, Databricks, Power BI, KPI reporting, segmentation, funnel analysis, ETL"),
        ("AI / ML", "Pandas, PySpark, PyTorch, TensorFlow, XGBoost, retrieval workflows, recommendation systems, prompt evaluation, LoRA"),
        ("Tools", "FastAPI, Flask, Streamlit, Docker, Git, GitHub, OpenAI API, REST APIs, Azure, joblib"),
    ]
    for label, content in lines:
        p = doc.add_paragraph(style="SkillLine")
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(content)


def add_footer_note(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "MASTER RESUME NOTE: This version keeps the full project bank so role-specific resumes can pull the strongest subset for data analyst, data scientist, AI, engineering, or fintech applications."
    )
    run.italic = True
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_resume() -> Path:
    doc = Document()
    style_document(doc)
    add_contact_block(doc)
    add_summary(doc)
    add_education(doc)
    add_experience(doc)
    add_projects(doc)
    add_research_and_skills(doc)
    add_footer_note(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_resume()
    print(path)
